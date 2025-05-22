import streamlit as st
from docx import Document
import pandas as pd
import json
import re
import io
import requests
from boxsdk import OAuth2, Client
from boxsdk.exception import BoxAPIException

st.set_page_config(page_title="Conga to Box Doc Gen", layout="centered")

# --- AI Model Configuration ---
ALL_MODELS_WITH_DESC = {
    "azure__openai__gpt_4_1_mini": "Azure OpenAI GPT-4.1 Mini: Lightweight multimodal model (Default for Box AI for Docs/Notes Q&A)",
    "google__gemini_2_0_flash_lite_preview": "Google Gemini 2.0 Flash Lite: Lightweight multimodal model (Preview)",
    "azure__openai__gpt_4o_mini": "Azure OpenAI GPT-4o Mini: Lightweight multimodal model",
    "azure__openai__gpt_4o": "Azure OpenAI GPT-4o: Highly efficient multimodal model for complex tasks",
    "azure__openai__gpt_4_1": "Azure OpenAI GPT-4.1: Highly efficient multimodal model for complex tasks",
    "azure__openai__gpt_o3": "Azure OpenAI GPT o3: Highly efficient multimodal model for complex tasks", 
    "azure__openai__gpt_o4-mini": "Azure OpenAI GPT o4-mini: Highly efficient multimodal model for complex tasks",
    "google__gemini_2_5_pro_preview": "Google Gemini 2.5 Pro: Optimal for high-volume, high-frequency tasks (Preview)",
    "google__gemini_2_5_flash_preview": "Google Gemini 2.5 Flash: Optimal for high-volume, high-frequency tasks (Preview)",
    "google__gemini_2_0_flash_001": "Google Gemini 2.0 Flash: Optimal for high-volume, high-frequency tasks",
    "google__gemini_1_5_flash_001": "Google Gemini 1.5 Flash: High volume tasks & latency-sensitive applications",
    "google__gemini_1_5_pro_001": "Google Gemini 1.5 Pro: Foundation model for various multimodal tasks",
    "aws__claude_3_haiku": "AWS Claude 3 Haiku: Tailored for various language tasks",
    "aws__claude_3_sonnet": "AWS Claude 3 Sonnet: Advanced language tasks, comprehension & context handling",
    "aws__claude_3_5_sonnet": "AWS Claude 3.5 Sonnet: Enhanced language understanding and generation",
    "aws__claude_3_7_sonnet": "AWS Claude 3.7 Sonnet: Enhanced language understanding and generation",
    "aws__titan_text_lite": "AWS Titan Text Lite: Advanced language processing, extensive contexts",
    "ibm__llama_3_2_instruct": "IBM Llama 3.2 Instruct: Instruction-tuned text model for dialogue, retrieval, summarization",
    "ibm__llama_3_2_90b_vision_instruct": "IBM Llama 3.2 90B Vision Instruct: Instruction-tuned vision model",
    "ibm__llama_4_scout": "IBM Llama 4 Scout: Natively multimodal model for text and multimodal experiences",
    "xai__grok_3_beta": "xAI Grok 3: Excels at data extraction, coding, summarization (Beta)",
    "xai__grok_3_mini_beta": "xAI Grok 3 Mini: Lightweight model for logic-based tasks (Beta)"
}
ALLOWED_MODEL_IDS = [
    "azure__openai__gpt_4o_mini", "azure__openai__gpt_4_1", "azure__openai__gpt_4_1_mini",
    "google__gemini_1_5_pro_001", "google__gemini_1_5_flash_001", "google__gemini_2_0_flash_001",
    "google__gemini_2_0_flash_lite_preview", "aws__claude_3_haiku", "aws__claude_3_sonnet",
    "aws__claude_3_5_sonnet", "aws__claude_3_7_sonnet", "aws__titan_text_lite",
    "ibm__llama_3_2_90b_vision_instruct", "ibm__llama_4_scout"
]
BOX_AI_MODELS_FOR_SELECTBOX = {"Default (Let Box Choose)": None}
for model_id_val in ALLOWED_MODEL_IDS:
    description = ALL_MODELS_WITH_DESC.get(model_id_val, "No description available")
    display_name = f"{model_id_val} ({description.split(':')[0]})" 
    BOX_AI_MODELS_FOR_SELECTBOX[display_name] = model_id_val
# --- End AI Model Configuration ---

# --- Function Definitions ---
def get_box_client():
    try:
        client_id = st.secrets["box"]["BOX_CLIENT_ID"]
        client_secret = st.secrets["box"]["BOX_CLIENT_SECRET"]
        developer_token = st.secrets["box"]["BOX_DEVELOPER_TOKEN"] 
        if not developer_token:
            st.sidebar.error("BOX_DEVELOPER_TOKEN is missing in Streamlit secrets.")
            return None 
        oauth = OAuth2(client_id=client_id, client_secret=client_secret, access_token=developer_token)
        client = Client(oauth)
        client.user().get() 
        return client
    except KeyError as e:
        st.sidebar.error(f"Box API credential missing: {e}.")
        return None
    except BoxAPIException as e:
        st.sidebar.error(f"Box API Error: {e.message} (Status: {e.status})")
        if e.status == 401: st.sidebar.warning("Expired/invalid Developer Token?")
        elif e.status == 503: st.sidebar.warning("Box API temporarily unavailable (503).")
        print(f"BoxAPIException: Status={e.status}, Code={e.code}, Msg={e.message}, ID={e.request_id}, URL={e.url}, Method={e.method}")
        return None
    except Exception as e:
        st.sidebar.error(f"Error connecting to Box: {str(e)}")
        return None

def get_folder_contents(client, folder_id):
    folder = client.folder(folder_id).get(fields=['name', 'parent'])
    items = client.folder(folder_id).get_items()
    return folder, items

def navigate_folders(client, current_folder_id):
    folder, items = get_folder_contents(client, current_folder_id)
    path_parts = []
    ancestor_tracer = folder
    while True:
        if not hasattr(ancestor_tracer, 'parent') or \
           not ancestor_tracer.parent or \
           ancestor_tracer.parent.id == "0":
            break 
        parent_ref = ancestor_tracer.parent
        try:
            parent_full_obj = client.folder(parent_ref.id).get(fields=["name", "parent"])
            path_parts.insert(0, {"id": parent_full_obj.id, "name": parent_full_obj.name})
            ancestor_tracer = parent_full_obj
        except Exception as e:
            st.warning(f"Could not get parent folder details for {parent_ref.id}: {e}")
            break
            
    breadcrumb_cols = st.columns(len(path_parts) + 1) 
    if breadcrumb_cols[0].button("Root", key="nav_root_button"):
        st.session_state.current_folder = "0"
        st.rerun()

    for i, part in enumerate(path_parts):
        if breadcrumb_cols[i+1].button(part["name"], key=f"nav_breadcrumb_{part['id']}"):
            st.session_state.current_folder = part["id"]
            st.rerun()
            
    st.subheader(f"Folder: {folder.name}")
    header_cols = st.columns([0.55, 0.15, 0.15, 0.15]) 
    header_cols[0].write("Name")
    header_cols[1].write("Type")
    header_cols[2].write("Select") 
    header_cols[3].write("Action")

    for item in items:
        item_cols = st.columns([0.55, 0.15, 0.15, 0.15])
        item_cols[0].write(f"{'📁' if item.type == 'folder' else '📄'} {item.name}")
        item_cols[1].write(item.type.capitalize())
        
        if item.type == "folder":
            if item_cols[3].button("Open", key=f"open_folder_{item.id}", help=f"Open folder {item.name}"):
                 st.session_state.current_folder = item.id
                 st.rerun()
        else: 
            file_extension = item.name.split('.')[-1].lower() if '.' in item.name else ''
            selected_as_template = st.session_state.get('template_file_id') == item.id
            selected_as_query = st.session_state.get('query_file_id') == item.id
            selected_as_schema = st.session_state.get('schema_file_id') == item.id
            checkbox_changed = False
            cb_key_suffix = item.id 
            if file_extension == 'docx':
                new_val_template = item_cols[2].checkbox("T", key=f"cb_template_{cb_key_suffix}", value=selected_as_template, help="Select as Template")
                if new_val_template != selected_as_template:
                    if new_val_template: st.session_state['template_file_id'] = item.id; st.session_state['template_file_name'] = item.name
                    else: st.session_state.pop('template_file_id', None); st.session_state.pop('template_file_name', None)
                    checkbox_changed = True
            elif file_extension in ['txt', 'csv']:
                new_val_query = item_cols[2].checkbox("Q", key=f"cb_query_{cb_key_suffix}", value=selected_as_query, help="Select as Query File")
                if new_val_query != selected_as_query:
                    if new_val_query: st.session_state['query_file_id'] = item.id; st.session_state['query_file_name'] = item.name; st.session_state['query_file_type'] = file_extension
                    else: st.session_state.pop('query_file_id', None); st.session_state.pop('query_file_name', None); st.session_state.pop('query_file_type', None)
                    checkbox_changed = True
            elif file_extension == 'json':
                new_val_schema = item_cols[2].checkbox("S", key=f"cb_schema_{cb_key_suffix}", value=selected_as_schema, help="Select as Schema File")
                if new_val_schema != selected_as_schema:
                    if new_val_schema: st.session_state['schema_file_id'] = item.id; st.session_state['schema_file_name'] = item.name
                    else: st.session_state.pop('schema_file_id', None); st.session_state.pop('schema_file_name', None)
                    checkbox_changed = True
            if checkbox_changed: st.rerun()

def download_box_file(client, file_id):
    return client.file(file_id).content()

def extract_merge_fields(docx_content):
    with io.BytesIO(docx_content) as docx_bytes:
        document = Document(docx_bytes)
        raw_fields = set()
        pattern = r"\{\{(.*?)\}\}" 
        for para in document.paragraphs:
            matches = re.findall(pattern, para.text)
            for match_text in matches: raw_fields.add(match_text.strip())
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para_in_cell in cell.paragraphs:
                        matches = re.findall(pattern, para_in_cell.text)
                        for match_text in matches: raw_fields.add(match_text.strip())
        
        cleaned_fields = {re.split(r'\s*(?:\\@|\||&)', f, maxsplit=1)[0].strip() for f in raw_fields}
        
        if len(raw_fields) != len(cleaned_fields) or any(r != c for r, c in zip(sorted(list(raw_fields)), sorted(list(cleaned_fields)))):
            st.info(f"Cleaned merge fields. Original: {len(raw_fields)}. Core: {len(cleaned_fields)}")
        return list(cleaned_fields)

def parse_conga_query(file_content, file_type):
    sobject_to_fields_map = {}
    all_unique_fields = set()
    all_unique_sobjects = set()

    def extract_from_soql_list(query_text_list):
        processed_queries_count = 0
        for query_text_item in query_text_list:
            query_text_item = str(query_text_item).strip()
            if not query_text_item: continue
            
            from_match = re.search(r"from\s+([a-zA-Z0-9_]+(?:__c|__r)?)\b", query_text_item, re.IGNORECASE)
            current_sobject_from_query = from_match.group(1) if from_match else None
            
            if current_sobject_from_query:
                all_unique_sobjects.add(current_sobject_from_query)
                sobject_to_fields_map.setdefault(current_sobject_from_query, set())

            if "select " in query_text_item.lower() and " from " in query_text_item.lower():
                processed_queries_count +=1
                select_match = re.search(r"select\s+(.+?)\s+from", query_text_item, re.IGNORECASE | re.DOTALL)
                if select_match:
                    fields_segment = select_match.group(1)
                    current_query_fields_list = re.findall(r"([\w\._()]+(?:\s+AS\s+[\w\._]+)?)(?:\s*,|\s+FROM)", fields_segment + " FROM", re.IGNORECASE | re.DOTALL)
                    cleaned_query_fields_set = {f.strip() for f in current_query_fields_list if f.strip()}
                    all_unique_fields.update(cleaned_query_fields_set)
                    if current_sobject_from_query: 
                        sobject_to_fields_map[current_sobject_from_query].update(cleaned_query_fields_set)
        return processed_queries_count

    queries_to_parse = []
    if file_type == 'csv':
        try:
            df_read = pd.read_csv(io.BytesIO(file_content), header=None, usecols=[0], skip_blank_lines=True) 
            if isinstance(df_read, pd.DataFrame):
                df_series = df_read.iloc[:, 0] if not df_read.empty else pd.Series(dtype=str)
            else: 
                df_series = df_read
            
            if not df_series.empty:
                queries_to_parse = df_series.dropna().astype(str).tolist()
            
            if queries_to_parse: st.info(f"CSV: Processing {len(queries_to_parse)} lines from first column as SOQL.")
            else:
                st.warning("CSV: No queries in first column after parsing. Using headers as fields if available.")
                df_headers = pd.read_csv(io.BytesIO(file_content), nrows=0) 
                all_unique_fields.update([str(col).strip() for col in df_headers.columns])
                if all_unique_fields: 
                    sobject_to_fields_map["UnknownFromCSVHeaders"] = set(all_unique_fields)
                    st.info(f"CSV (Fallback): Using column headers as query fields: {list(all_unique_fields)}")
                return {k: sorted(list(v)) for k,v in sobject_to_fields_map.items()}, all_unique_fields, all_unique_sobjects
        except Exception as e:
            st.error(f"Error parsing CSV: {e}. Trying as TXT."); return parse_conga_query(file_content, 'txt')
    elif file_type == 'txt':
        text_content = file_content.decode("utf-8")
        queries_to_parse = [q.strip() for q in re.split(r';\s*\n?|\n\s{2,}\n?|\n\s*$', text_content, flags=re.MULTILINE) if q.strip()]
        if queries_to_parse: st.info(f"TXT: Found {len(queries_to_parse)} potential SOQL.")
    else: st.error(f"Unsupported query file type: {file_type}"); return {}, set(), set()

    processed_count = extract_from_soql_list(queries_to_parse) if queries_to_parse else 0
    msg_type = file_type.upper()
    if processed_count > 0: st.success(f"{msg_type}: Processed {processed_count} SOQL. Fields: {len(all_unique_fields)}, SObjects: {len(all_unique_sobjects)}.")
    elif queries_to_parse: st.warning(f"{msg_type}: No SOQL processed from {len(queries_to_parse)} candidates.")
    else: st.warning(f"{msg_type}: No query strings found.")
    return {k: sorted(list(v)) for k, v in sobject_to_fields_map.items()}, all_unique_fields, all_unique_sobjects

def generate_prompt_single_call(merge_fields, sobject_to_fields_map, all_query_sobjects):
    merge_fields_str_list = []
    for mf in sorted(list(set(merge_fields))):
        mf_lower = mf.lower()
        possible_sources = []
        for sobj in all_query_sobjects: 
            sobj_clean_prefix = sobj.lower().replace("__c", "").replace("__r", "") + "_"
            if mf_lower.startswith(sobj_clean_prefix):
                if sobj not in [s.split(" ")[0] for s in possible_sources]: possible_sources.append(f"{sobj} (prefix match)"); break 
        if not possible_sources: 
            for sobj, fields in sobject_to_fields_map.items():
                for q_field in fields:
                    q_field_core = q_field.split('.')[-1].lower().replace("__c", "").replace("__r", "")
                    mf_core = mf_lower.replace("__c", "").replace("__pc","") 
                    if mf_core == q_field_core or mf_core in q_field_core or q_field_core in mf_core :
                        if sobj not in [s.split(" ")[0] for s in possible_sources]: possible_sources.append(f"{sobj} (related query field: {q_field})"); break 
                if len(possible_sources) > 1: break 
        hint = f" (Hint: Likely from SObject(s) - {'; '.join(possible_sources)[:200]})" if possible_sources else " (Hint: general search in schema needed)"
        merge_fields_str_list.append(f"- {mf}{hint}")

    detailed_merge_fields_list = "\n".join(merge_fields_str_list)
    all_query_sobjects_str = ", ".join(sorted(list(all_query_sobjects))) if all_query_sobjects else "Not specified, use full schema"

    prompt = (
        "You are an AI assistant performing a precise field mapping task from a Conga document generation system to a Salesforce schema fields for Box Doc Gen. "
        "You have been provided with a Salesforce schema file as context (see the 'items' passed to the API call). Use this schema file as the PRIMARY and SOLE source of truth for valid Salesforce field paths and their data types.\n\n"
        "Conga Template Merge Fields to Map (each with a hint about potential source SObjects based on associated Conga queries):\n"
        f"{detailed_merge_fields_list}\n\n"
        "Primary Salesforce SObjects involved in the source Conga queries (consult the provided schema file for their detailed structure and relationships):\n"
        f"{all_query_sobjects_str}\n\n"
        "CRITICAL INSTRUCTIONS FOR ACCURACY AND FORMATTING:\n"
        "1. Schema is Truth: ALL 'BoxField' paths MUST EXACTLY match paths found in the provided Salesforce schema file. Do NOT invent fields, relationships, or suffixes (like adding '_c' if it's not in the schema for that field).\n"
        "2. Custom Fields: Salesforce custom fields typically end in '__c'. If a CongaField ends in '__pc' (e.g., MyField__pc), search the schema for 'SObject.MyField__pc' first. If not found, then search for 'SObject.MyField__c'.\n"
        "3. Person Accounts: If the schema indicates Person Account fields (e.g., Account.PersonEmail, Account.PersonBirthdate - these are standard fields, not custom), map 'ACCOUNT_PERSON...' CongaFields accordingly. These standard Person Account fields do NOT end in '__c'.\n"
        "4. User Fields: CongaFields prefixed 'USER_' (e.g., USER_CITY) should generally map to fields on the '$User' global object (e.g., '$User.City') IF such a field exists in the schema. If context implies a user from a related record (e.g., 'Matter_Team_Members__c.User__r'), use that full valid path from the schema.\n"
        "5. System Fields: For fields like 'Today', use system variables like '$System.Today' for Date or DateTime if present/inferable.\n"
        "6. No Invention/Guessing: If you find a field in the schema that is a *close but not exact* match for a CongaField, and you are not highly confident, it is BETTER to leave the BoxField and FieldType BLANK for that CongaField. DO NOT output conditional text like '(if exists)' or notes within the BoxField/FieldType columns.\n\n"
        "TASK:\n"
        "For EACH 'Conga Template Merge Field' from the list above, provide its corresponding Salesforce field path from the schema file. "
        "Use the hints and SObject list to guide your search within the schema. After identifying a potential BoxField, CRITICALLY RE-VERIFY that this exact path exists in the provided schema file. If it does not, the BoxField must be left blank.\n"
        "STRICT OUTPUT FORMAT: Your entire response MUST BE ONLY VALID CSV data. "
        "1. Start with the exact header row: CongaField,BoxField,FieldType\n"
        "2. Each subsequent line is a data row. Each data row MUST correspond to one CongaField from the input list.\n"
        "3. Each data row MUST have exactly three values, separated by a SINGLE comma.\n"
        "   - Correct Mapped Row Example: ActualCongaField,Account.Actual_Salesforce_Field__c,Text\n"
        "   - Correct Unmapped Row Example: UnmappedCongaField,,\n"
        "4. NO EXTRA TEXT: Do NOT include ANY text, notes, explanations, apologies, or conversational remarks before, after, or within the CSV data cells.\n"
        "5. DELIMITERS: Use ONLY a single comma as a delimiter. Do not use tabs or multiple spaces.\n"
        "6. QUOTES: Do not add quotes unless a value itself contains a comma that needs escaping (standard CSV behavior).\n"
        "Failure to follow this CSV format will render the output unusable. Each CongaField from the input list must appear exactly once in your CSV output.\n\n"
        "CSV Columns Definition:\n"
        "1. CongaField: The exact Conga merge field (from the list above, without its hint part).\n"
        "2. BoxField: The full, valid path from the Salesforce Schema file. Must be verifiable in the schema.\n"
        "3. FieldType: The data type (e.g., Text, Number, Date, Boolean, Picklist, Id, RichText, Lookup) from the schema for the BoxField.\n"
        "If a CongaField cannot be confidently mapped to an existing field in the provided schema, leave its BoxField and FieldType columns BLANK.\n"
        "Example Row: Template_Account_Name,Account.Name,Text"
    )
    return prompt

def call_box_ai(prompt, grounding_file_id, developer_token, model_id=None, temperature=0.2): # Default temperature
    url = "https://api.box.com/2.0/ai/text_gen"
    headers = {"Authorization": f"Bearer {developer_token}", "Content-Type": "application/json"}
    items_payload = [{"type": "file", "id": str(grounding_file_id), "content_type": "text_content"}] if grounding_file_id else []
    
    data = {"prompt": prompt, "items": items_payload}
    model_used_msg = "Using default Box AI model."
    if model_id: 
        data["model"] = model_id
        model_used_msg = f"Using Box AI model: {model_id}"
    if temperature is not None: 
        data["temperature"] = temperature 
        model_used_msg += f" with temperature: {temperature}"
    
    st.info(model_used_msg); 
    displayed_data = data.copy()
    if len(json.dumps(displayed_data.get("prompt", ""))) > 1000: # Truncate long prompts in UI
        displayed_data["prompt"] = displayed_data["prompt"][:1000] + "...\n(Prompt truncated in UI display for brevity)"
    st.info("Box AI Request Payload (see console for full details if large):"); st.json(displayed_data) 

    print(f"--- Sending Box AI Request ---\nURL: {url}\nModel: {data.get('model', 'Default')}\nTemperature: {data.get('temperature', 'Not Set')}\nPayload Len: {len(json.dumps(data))}\n-----------------------------")
    response = requests.post(url, headers=headers, json=data)
    print(f"--- Received Box AI Response (Status: {response.status_code}) ---")
    resp_text = None
    try: 
        resp_data = response.json(); print(f"Response JSON (first 500 chars): {json.dumps(resp_data, indent=2)[:500]}...")
        if response.ok: resp_text = resp_data.get("answer", "") 
    except json.JSONDecodeError: print(f"Response (Non-JSON, first 500): {response.text[:500]}..."); resp_text = response.text if response.ok else None
    except Exception as e: print(f"Error processing response: {e}\nRaw (first 500): {response.text[:500]}...")
    print("-----------------------------")
    if response.ok: return resp_text
    else: st.error(f"Box AI request failed: {response.status_code} — {response.text}"); return None

def convert_response_to_df(text):
    if not text or not isinstance(text, str): st.warning("AI response empty/invalid."); return pd.DataFrame()
    csv_match = re.search(r"```(?:csv)?\s*([\s\S]*?)\s*```", text, re.IGNORECASE)
    csv_text = csv_match.group(1).strip() if csv_match else text.strip()
    if csv_match: st.info("Extracted CSV block from AI response.")
    else: st.info("No CSV markdown block found, parsing response as is.")
    lines = csv_text.strip().splitlines()
    header_idx = -1; expected_hdrs_check = ["CongaField", "BoxField"]
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if all(eh.lower() in line_lower for eh in expected_hdrs_check) and \
           not any(mf_prefix.lower() in line_lower for mf_prefix in ["account_", "case_", "contact_", "user_", "$system", "today"]):
            header_idx = i; break
    final_cols = ["CongaField", "BoxField", "FieldType"]
    if header_idx == -1:
        st.warning(f"CSV header (CongaField, BoxField) not reliably found.")
        for i, line in enumerate(lines): 
            if line.strip() and not re.match(r"^\s*\(Note:|\*\s|---|^\s*Here are|^\s*Please note|^\s*This CSV", line, re.IGNORECASE):
                header_idx = i; st.info(f"Using line {i+1} as potential header: '{line}'"); break
        if header_idx == -1: 
            st.error("No plausible header line found."); st.text_area("Processed AI Text", csv_text, height=100)
            return pd.DataFrame(columns=final_cols)

    header_from_ai_raw = [h.strip() for h in lines[header_idx].split(",")]
    if len(header_from_ai_raw) < len(final_cols): header_from_ai_raw.extend([""] * (len(final_cols) - len(header_from_ai_raw)))
    elif len(header_from_ai_raw) > len(final_cols): header_from_ai_raw = header_from_ai_raw[:len(final_cols)]

    data_list = []
    for line_num, line_content in enumerate(lines[header_idx+1:]):
        line_content_s = line_content.strip()
        if not line_content_s: continue
        if re.match(r"^\s*\(Note:|\*\s|---|^\s*Here are|^\s*Please note|^\s*This CSV|^\s*Example Row:", line_content_s, re.IGNORECASE) or \
           (not line_content_s[0].isalnum() and line_content_s[0] not in ['{', '"', '(', '$'] and line_content_s.count(',') == 0 and len(line_content_s) < 20) : # Added '$' for system vars
            if line_content_s: st.info(f"Skipping non-data line: '{line_content_s}'"); continue
        
        split_vals = [v.strip() for v in line_content_s.split(",")]
        if len(split_vals) == 1 and line_content_s.count(',') == 0 and len(re.split(r'\s{2,}|\\t', line_content_s)) >= 2 :
            alt_split_vals = re.split(r'\s{2,}|\\t', line_content_s) 
            if len(alt_split_vals) >=2 and len(alt_split_vals) <=len(final_cols):
                st.info(f"Line '{line_content_s}' has few commas, using multi-space/tab split -> {alt_split_vals}")
                split_vals = alt_split_vals
        
        row_data_ordered = [""] * len(final_cols)
        for i_val in range(len(final_cols)): # Iterate up to the number of final columns
            if i_val < len(split_vals):
                # If this is the last column we expect AND there are more values in split_vals, join them
                if i_val == len(final_cols) - 1 and len(split_vals) > len(final_cols):
                    row_data_ordered[i_val] = ",".join(split_vals[i_val:])
                else:
                    row_data_ordered[i_val] = split_vals[i_val]
        
        if row_data_ordered[0]: data_list.append(dict(zip(final_cols, row_data_ordered)))
        elif any(val for val in row_data_ordered): st.info(f"Skipping row with empty CongaField: {row_data_ordered}")
            
    if not data_list: st.warning("No valid data rows extracted."); return pd.DataFrame(columns=final_cols)
    return pd.DataFrame(data_list, columns=final_cols)

# --- Streamlit App UI & Main Logic ---
if 'current_folder' not in st.session_state: st.session_state['current_folder'] = "0"
st.title("Conga Template to Box Doc Gen Mapper")

with st.sidebar:
    st.header("Box File Browser")
    client = get_box_client() 
    if client: 
        navigate_folders(client, st.session_state['current_folder']) 
        st.subheader("Selected Files")
        st.write(f"Template: {st.session_state.get('template_file_name', 'Not selected')}")
        st.write(f"Query: {st.session_state.get('query_file_name', 'Not selected')}")
        st.write(f"Schema: {st.session_state.get('schema_file_name', 'Not selected')}")
        st.divider(); st.subheader("AI Configuration")
        selected_model_display_name = st.selectbox("Choose Box AI Model:", options=list(BOX_AI_MODELS_FOR_SELECTBOX.keys()), index=0, key="ai_model_selector")
        ai_temperature = st.slider("AI Temperature (0.0=Deterministic, 1.0=Creative):", min_value=0.0, max_value=1.0, value=0.2, step=0.05, key="ai_temperature_slider")

    else: st.warning("Box client not initialized. File browsing/AI unavailable.")

if st.button("Generate Field Mapping", key="generate_mapping_button"):
    if 'client' not in locals() or client is None: client = get_box_client() 
    if client is None: st.error("Box client init failed. Check sidebar errors."); st.stop()
    if not all(st.session_state.get(k) for k in ['template_file_id', 'query_file_id', 'schema_file_id']):
        st.warning("Please select all DOCX, Query, and Schema files."); st.stop()
    
    chosen_model_display_name = st.session_state.ai_model_selector 
    final_selected_model_id = BOX_AI_MODELS_FOR_SELECTBOX[chosen_model_display_name]
    final_ai_temperature = st.session_state.ai_temperature_slider


    st.info("Processing files and generating mapping...")
    prog_bar = st.progress(0.0, text="Initializing...")
    try:
        prog_bar.progress(0.1, text="Downloading files...");
        template_content = download_box_file(client, st.session_state['template_file_id'])
        query_content = download_box_file(client, st.session_state['query_file_id'])
        prog_bar.progress(0.2, text="Extracting merge fields...");
        merge_fields = extract_merge_fields(template_content)
        if merge_fields: st.success(f"Extracted {len(merge_fields)} merge fields.")
        else: st.warning("No merge fields found.")
        prog_bar.progress(0.3, text="Parsing Conga query...");
        sobject_to_fields_map, _, all_query_sobjects = parse_conga_query(query_content, st.session_state['query_file_type'])
        prog_bar.progress(0.4, text="Validating schema file ID...");
        schema_file_id = st.session_state.get('schema_file_id')
        if not schema_file_id: st.error("Schema file ID missing for AI call."); st.stop()
        st.success(f"Schema file ID {schema_file_id} will be used for AI context.")
        if not merge_fields: st.error("No merge fields to map."); st.stop()

        prog_bar.progress(0.5, text="Calling Box AI...");
        prompt = generate_prompt_single_call(merge_fields or [], sobject_to_fields_map, all_query_sobjects or [])
        ai_resp = call_box_ai(
            prompt, 
            schema_file_id, 
            st.secrets["box"]["BOX_DEVELOPER_TOKEN"], 
            final_selected_model_id,
            temperature=final_ai_temperature
        )
        prog_bar.progress(0.8, text="Processing AI response...");
        df_map = pd.DataFrame()
        if ai_resp: df_map = convert_response_to_df(ai_resp) 
        else: st.warning("No response text from Box AI.")
        prog_bar.progress(1.0, text="Completed!"); prog_bar.empty()

        if not df_map.empty:
            exp_cols = ["CongaField", "BoxField", "FieldType"]
            for col in exp_cols: 
                if col not in df_map.columns: df_map[col] = "" 
            df_map = df_map[exp_cols] 
            df_map.drop_duplicates(subset=['CongaField', 'BoxField'], inplace=True, keep='first')
            st.subheader("Generated Field Mapping"); st.dataframe(df_map)
            csv_buf = io.StringIO(); df_map.to_csv(csv_buf, index=False)
            st.download_button("Download Mapping CSV", csv_buf.getvalue(), "field_mapping.csv", "text/csv")
        else: st.warning("No valid field mappings generated by Box AI.")
    except Exception as e:
        st.error(f"An error occurred during processing: {str(e)}"); import traceback; st.text(traceback.format_exc())
        if 'prog_bar' in locals() and prog_bar is not None : prog_bar.empty()
